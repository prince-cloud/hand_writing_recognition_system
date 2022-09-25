import secrets
from django.http import HttpRequest, JsonResponse
from django.shortcuts import render, redirect
from django.views.decorators.csrf import csrf_exempt
from django.contrib import messages
# packages
from google.cloud import vision_v1, translate_v2, texttospeech_v1
from django.conf import settings
import base64
from django.utils import timezone
from services.forms import ScanTextForm, TranslateTextForm, TextToSpeechForm
from django.template.loader import get_template
from xhtml2pdf import pisa
from io import BytesIO
from django.http import HttpResponse
from docx import Document
from io import StringIO
#pdf = render_to_pdf('pdf_template.html', {'purchase': order})
# return HttpResponse(pdf, content_type='application/pdf')

# Create your views here.
def render_to_pdf(template_src, context_dict={}):
    template = get_template(template_src)
    html = template.render(context_dict)
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode("ISO-8859-1")), result)
    if not pdf.err:
        return HttpResponse(result.getvalue(), content_type='application/pdf')
    return None

def index(request):
    return render(request, "index.html", {"ScanTextForm": ScanTextForm()})


def englishTwi(request: HttpRequest):
    pass


def exportScanText(request):
    text = request.GET.get('text')
    ex_type = request.GET.get('ex_type')
    if ex_type == 'word':
       document = Document()
       if text:
            document.add_paragraph(text)
            f = BytesIO()
            document.save(f)
            length = f.tell()
            f.seek(0)
            response = HttpResponse(
                f.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = 'attachment; filename=' + 'word_document.docx'
            response['Content-Length'] = length
            return response
    else:
        if text:
            pdf = render_to_pdf('pdf_template.html', {'text': text})
            return HttpResponse(pdf, content_type='application/pdf')
    return render(request, 'pages/scan_text.html', {'text':text})


def exportTranslateText(request):
    text = request.GET.get('text')
    ex_type = request.GET.get('ex_type')
    if ex_type == 'word':
       document = Document()
       if text:
            document.add_paragraph(text)
            f = BytesIO()
            document.save(f)
            length = f.tell()
            f.seek(0)
            response = HttpResponse(
                f.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = 'attachment; filename=' + 'word_document.docx'
            response['Content-Length'] = length
            return response
    else:
        if text:
            pdf = render_to_pdf('pdf_template.html', {'text': text})
            return HttpResponse(pdf, content_type='application/pdf')
    return render(request, 'pages/translate_text.html', {'text':text})



def scan_text(request: HttpRequest):
    docText = ""
    b64_img = ""
    myimageurl = ""
    print(request.headers)
    if request.method == "POST":
        form = ScanTextForm(data=request.POST, files=request.FILES)
        if form.is_valid():
            data = form.cleaned_data
            scan_form = form.save(commit=False)
            image_file = data["image"]
            client = vision_v1.ImageAnnotatorClient()
            content = image_file.read()
            image = vision_v1.types.Image(content=content)
            response = client.document_text_detection(image=image)
            docText = response.full_text_annotation.text
            print(docText)
            scan_form.save()
            myimageurl = "http://localhost:8000/media/scanned_images/" + str(image_file)

        else:
            messages.error(request, "form invalid")
    else:
        form = ScanTextForm()

    return render(
        request,
        "pages/scan_text.html",
        {
            "form": form,
            "docText": docText,
            # "translate_form": TranslateTextForm(),
            "selected_image": b64_img,
            "myimageurl": myimageurl,
        },
    )


def translateText(request):
    translation = ""
    translate_to = ""
    translate_client = translate_v2.Client()

    if request.method == "POST":
        form = TranslateTextForm(data=request.POST, files=request.FILES)
        if form.is_valid():
            data = form.cleaned_data
            translateForm = form.save(commit=False)
            translateForm.text = data["text"]
            translateForm.translate_to = data["translate_to"]
            text = data["text"]
            translate_to = data["translate_to"]

            # translator = Translator(to_lang=f"{translate_to}")
            # translation = translator.translate(f"{text}")
            translate_text = translate_client.translate(
                text,
                target_language=translate_to,
            )

            print(translation)
            translateForm.translated_text = translate_text["translatedText"]
            translation = translate_text["translatedText"]
            translateForm.save()
        else:
            messages.error(request, "invalid data entry")
    else:
        form = TranslateTextForm()

    return render(
        request,
        "pages/translate_text.html",
        {"form": form, "translation": translation, "translate_to": translate_to},
    )


def generate_file_name():
    now = timezone.now()
    token = secrets.token_urlsafe(5)
    media_root = settings.MEDIA_ROOT
    output = f"/audio/{now.date()}-{now.hour}-{now.minute}-{token}.mp3"
    return output


def texttospeech(request):
    client = texttospeech_v1.TextToSpeechClient()
    filename = ""
    if request.method == "POST":
        form = TextToSpeechForm(data=request.POST, files=request.FILES)
        if form.is_valid():
            data = form.cleaned_data
            ttsform = form.save(commit=False)
            text = data["text"]

            text = f"<speak> {text} </speak>"

            synthesis_input = texttospeech_v1.SynthesisInput(ssml=text)

            voice = texttospeech_v1.VoiceSelectionParams(
                language_code="en-US",
                ssml_gender=texttospeech_v1.SsmlVoiceGender.MALE,
            )
            audio_config = texttospeech_v1.AudioConfig(
                audio_encoding=texttospeech_v1.AudioEncoding.MP3
            )
            response = client.synthesize_speech(
                input=synthesis_input,
                voice=voice,
                audio_config=audio_config,
            )
            print("======== printing response: ", response)
            filename = generate_file_name()
            print(settings.MEDIA_ROOT, filename)
            save_destination = str(settings.MEDIA_ROOT) + filename
            with open(save_destination, "wb") as output:
                output.write(response.audio_content)

            ttsform.save()
        else:
            print("invalid data entry")
            messages.error(request, "invalid data entry")
    else:
        form = TextToSpeechForm()

    return render(
        request,
        "pages/tts.html",
        {"form": form, "filename": settings.MEDIA_URL + filename},
    )


@csrf_exempt
def text_to_speech_api(request: HttpRequest):
    success = True
    message = ""
    client = texttospeech_v1.TextToSpeechClient()
    filename = ""
    if request.method == "POST":
        form = TextToSpeechForm(data=request.POST, files=request.FILES)
        if form.is_valid():
            data = form.cleaned_data
            ttsform = form.save(commit=False)
            text = data["text"]

            text = f"<speak> {text} </speak>"

            synthesis_input = texttospeech_v1.SynthesisInput(ssml=text)

            voice = texttospeech_v1.VoiceSelectionParams(
                language_code="en-US",
                ssml_gender=texttospeech_v1.SsmlVoiceGender.MALE,
            )
            audio_config = texttospeech_v1.AudioConfig(
                audio_encoding=texttospeech_v1.AudioEncoding.MP3
            )
            response = client.synthesize_speech(
                input=synthesis_input,
                voice=voice,
                audio_config=audio_config,
            )
            filename = generate_file_name()
            print(settings.MEDIA_ROOT, filename)
            save_destination = str(settings.MEDIA_ROOT) + filename
            with open(save_destination, "wb") as output:
                output.write(response.audio_content)

            ttsform.save()
            success = True
            message = "Text converted successfully"
        else:
            print("invalid data entry")
            success = False
            message = "Failed to convert text to speech"
    else:
        success = False
        message = "Only Post requests allowed"

    return JsonResponse(
        {
            "success": success,
            "message": message,
            "filename": (settings.MEDIA_URL + filename).replace("//", "/"),
            #"filename": ('https://585f-102-176-94-120.eu.ngrok.io' + filename).replace("//", "/"),

        }
    )


@csrf_exempt
def scan_text_api(request: HttpRequest):
    docText = ""
    if request.method == "POST":
        form = ScanTextForm(data=request.POST, files=request.FILES)
        if form.is_valid():
            data = form.cleaned_data
            scan_form = form.save(commit=False)
            image_file = data["image"]
            client = vision_v1.ImageAnnotatorClient()
            content = image_file.read()
            image = vision_v1.types.Image(content=content)
            response = client.document_text_detection(image=image)
            docText = response.full_text_annotation.text
            print(docText)
            scan_form.save()

        else:
            return JsonResponse(
                {
                    "docText": docText,
                    "success": False,
                    "message": "Translation faild. Please try again later.",
                }
            )
    else:
        return JsonResponse(
            {
                "docText": docText,
                "success": False,
                "message": "Translation faild. Please try again later.",
            }
        )
    return JsonResponse(
        {"docText": docText, "success": True, "message": "Translated successfully"}
    )


@csrf_exempt
def translateText_api(request: HttpRequest):
    translation = ""
    translate_to = ""
    translate_client = translate_v2.Client()

    if request.method == "POST":
        form = TranslateTextForm(data=request.POST, files=request.FILES)
        if form.is_valid():
            data = form.cleaned_data
            translateForm = form.save(commit=False)
            translateForm.text = data["text"]
            translateForm.translate_to = data["translate_to"]

            text = data["text"]
            translate_to = data["translate_to"]

            # translator = Translator(to_lang=f"{translate_to}")
            # translation = translator.translate(f"{text}")
            translation_text = translate_client.translate(
                text,
                target_language=translate_to,
            )

            # print (translation['translatedText'])
            translateForm.translated_text = translation_text["translatedText"]
            translation = translation_text["translatedText"]
            translateForm.save()
        else:
            return JsonResponse(
                {
                    "translation": translation,
                    "translated_to": translate_to,
                    "success": False,
                    "message": "Translation faild. Please try again later.",
                }
            )

    else:
        return JsonResponse(
            {
                "translation": translation,
                "translated_to": translate_to,
                "success": False,
                "message": "form failed",
            }
        )

    return JsonResponse(
        {
            "translation": translation,
            "translated_to": translate_to,
            "success": True,
            "message": "translation success",
        }
    )



#def exportToWord(request, text):
#    #text = request.GET.get('text')
#    document = Document()
#    if text:
#        document.add_paragraph(text)
#        f = BytesIO()
#        document.save(f)
#        length = f.tell()
#        f.seek(0)
#        response = HttpResponse(
#            f.getvalue(),
#            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#        )
#        response['Content-Disposition'] = 'attachment; filename=' + 'word_document.docx'
#        response['Content-Length'] = length
#        return response
#    return render(request, 'pages/scan_text.html', {'text':text})
#
#
#def exportToPdf(request):
#    text = request.GET.get('text')
#    if text:
#        pdf = render_to_pdf('pdf_template.html', {'text': text})
#        return HttpResponse(pdf, content_type='application/pdf')
#    return render(request, 'pages/scan_text.html', {'text':text})
#